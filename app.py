import os
from flask import Flask, render_template, request, redirect, send_file
from werkzeug.utils import secure_filename

from io import BytesIO
from docx import Document

# Import de la fonction main depuis transcriptor.py
from model import *

app = Flask(__name__)

# Configuration pour l'upload de fichiers
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'wav', 'mp3', 'ogg'}  # Types de fichiers autorisés

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Vérifie si le fichier a une extension autorisée
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Route pour la page d'accueil
@app.route("/")
def index():
    return render_template("index.html")

# Route pour gérer le fichier uploadé
@app.route("/upload", methods=["POST"])
def upload_file():
    if 'file' not in request.files:
        return redirect(request.url)
    
    file = request.files['file']
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        # Appel de la fonction main du backend avec le fichier uploadé
        output_data = run_model(file_path)

        # Retourne le résultat dans le template
        return render_template("index.html", output_data=output_data)

    return redirect(request.url)

@app.route("/save_txt", methods=["POST"])
def save_txt_file():
    output_data = request.form['output_data']
    
    # Créer un fichier texte
    txt_io = BytesIO()
    txt_io.write(output_data.encode('utf-8'))  # Écrire les données dans le fichier texte
    txt_io.seek(0)

    # Envoyer le fichier en réponse
    return send_file(txt_io, as_attachment=True, download_name="prediction.txt", mimetype="text/plain")

@app.route("/save", methods=["POST"])
def save_file():
    output_data = request.form['output_data']
    
    # Créer un document Word
    doc = Document()
    doc.add_heading('Prédiction', level=1)
    doc.add_paragraph(output_data)

    # Sauvegarder le document dans un objet BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    # Envoyer le fichier en réponse
    return send_file(doc_io, as_attachment=True, download_name="prediction.doc", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
